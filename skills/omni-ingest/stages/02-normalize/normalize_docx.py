"""
normalize_docx.py — omni-ingest Stage 2 normalizer: Word .docx → Markdown.

Uses python-docx to extract full document text preserving paragraph structure.
Extracts core_properties (title, author, created, modified).

API:
  def normalize_docx(
      docx_path: str | Path,
      workflow_id: str,
      original_filename: str | None = None,
  ) -> tuple[str, dict]:
      # Returns (markdown_content, meta_dict)
      # Writes bus/content.<workflow_id>.md + bus/content.<workflow_id>.meta.json

Run: python -m pytest tests/test_normalize_docx.py -v
"""

from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document  # python-docx

# Resolve relative to this file's location
_SKILL_ROOT = Path(__file__).parent.parent.parent
_BUS_DIR = _SKILL_ROOT / "bus"


def _extract_core_properties(doc: Document) -> dict:
    """Extract docx core properties if available."""
    try:
        cp = doc.core_properties
        return {
            "title": cp.title or "",
            "author": cp.author or "",
            "subject": cp.subject or "",
            "keywords": cp.keywords or "",
            "created": cp.created.isoformat() if cp.created else "",
            "modified": cp.modified.isoformat() if cp.modified else "",
            "last_modified_by": cp.last_modified_by or "",
        }
    except Exception:
        return {}


def _docx_to_markdown(doc: Document) -> str:
    """Convert python-docx paragraphs + tables to Markdown."""
    parts: list[str] = []

    for element in doc.element.body:
        if element.tag.endswith("p"):
            # It's a paragraph
            para = next((p for p in doc.paragraphs if p._element is element), None)
            if para is None:
                continue
            text = para.text.strip()
            if not text:
                parts.append("")
                continue
            # Detect heading by style
            style_name = para.style.name.lower() if para.style else ""
            if "heading 1" in style_name or "title" in style_name:
                parts.append(f"# {text}")
            elif "heading 2" in style_name:
                parts.append(f"## {text}")
            elif "heading 3" in style_name:
                parts.append(f"### {text}")
            else:
                parts.append(text)
        elif element.tag.endswith("tbl"):
            # It's a table
            tbl = next((t for t in doc.tables if t._element is element), None)
            if tbl:
                parts.append(_table_to_markdown(tbl))

    # Remove more than 2 consecutive blank lines
    lines = "\n".join(parts).splitlines()
    collapsed = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 1:
                collapsed.append(line)
        else:
            blank_count = 0
            collapsed.append(line)
    while collapsed and collapsed[-1] == "":
        collapsed.pop()
    return "\n".join(collapsed) + "\n"


def _table_to_markdown(table) -> str:
    """Convert a python-docx table to a Markdown table."""
    rows = table.rows
    if not rows:
        return ""
    # Header row
    header = [cell.text.strip() for cell in rows[0].cells]
    sep = ["---"] * len(header)
    body = [[cell.text.strip() for cell in row.cells] for row in rows[1:]]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def normalize_docx(
    docx_path: str | Path,
    workflow_id: str,
    original_filename: str | None = None,
) -> tuple[str, dict]:
    """
    Convert a .docx file to Markdown and write to bus/.

    Parameters
    ----------
    docx_path : str | Path
        Path to the .docx file.
    workflow_id : str
        Workflow identifier.
    original_filename : str | None
        Original filename for metadata.

    Returns
    -------
    (markdown_content, meta_dict) tuple.

    Raises
    ------
    FileNotFoundError
        If the .docx file does not exist.
    ValueError
        If the file is not a valid .docx.
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"docx not found: {docx_path}")
    if not path.suffix.lower() == ".docx":
        raise ValueError(f"not a .docx file: {docx_path}")

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"failed to open .docx: {exc}") from exc

    core = _extract_core_properties(doc)
    markdown_content = _docx_to_markdown(doc)

    # Build meta
    meta = {
        "originalFilename": original_filename or path.name,
        "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "normalizedMimeType": "text/markdown",
        "title": core.get("title", ""),
        "author": core.get("author", ""),
        "subject": core.get("subject", ""),
        "keywords": core.get("keywords", ""),
        "created": core.get("created", ""),
        "modified": core.get("modified", ""),
        "normalizedAt": datetime.now(timezone.utc).isoformat(),
        "lineCount": len(markdown_content.splitlines()),
        "wordCount": len(markdown_content.split()),
    }

    # Write bus files
    content_path = _BUS_DIR / f"content.{workflow_id}.md"
    meta_path = _BUS_DIR / f"content.{workflow_id}.meta.json"
    content_path.parent.mkdir(parents=True, exist_ok=True)
    content_path.write_text(markdown_content, encoding="utf-8")
    meta_path.write_text(json.dumps(meta, indent=2), encoding="utf-8")

    return markdown_content, meta


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python normalize_docx.py <file.docx> <workflow_id>")
        sys.exit(1)
    docx_path, workflow_id = sys.argv[1], sys.argv[2]
    content, meta = normalize_docx(docx_path, workflow_id)
    print(f"Converted: {meta['title'] or meta['originalFilename']}")
    print(json.dumps(meta, indent=2))
